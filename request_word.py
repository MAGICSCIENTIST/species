import unicodedata
from docx import Document
from docx.text.paragraph import Paragraph
import re

from bs4 import BeautifulSoup as Soup
import urllib

import utils

Paragraph.text = property(lambda self: GetParagraphText(self))
Paragraph.link = property(lambda self: GetLinkUrl(self))

class ImageObject:
    def __init__(self, url, name, isDownloaded=False):
        self.url = url
        self.name = name
        self.isDownloaded = isDownloaded

class DetailObject:
    def __init__(self, properties):
        self.images = []
        self.properties = properties

    def setDescribe(self, describe):
        self.describe = describe

    def setSpecieCategory(self, specieCategory):
        self.specieCategory = specieCategory

    def addImage(self, image: ImageObject):
        self.images.append(image)

class Node:
    def __init__(self, name, properties):
        self.name = name        
        self.properties = properties
        self.dataSource = None
        self.describe = None
        self.children = []
    def __init__(self, properties, dataSource, nameField="中文名"):
        self.properties = properties
        self.dataSource = dataSource
        self.describe = None
        self.children = []
        for key in properties.keys():
            if key == nameField:
                self.name = properties[key]
                break
            else:
                self.properties[key] = dataSource[properties[key]]
        


    def addChild(self, node) :
        self.children.append(node)

    def addDescribe(self, describe: DetailObject):
        self.describe = describe

    def get(self, key):
        return self.properties[key]
    
    def set(self, key, value):
        self.properties[key] = value

def createGenerator(document, properties_words, nameField="中文名"):
    parentChain = []
    for tb in  document.tables:
        for i, row in enumerate(tb.rows):                        
            if i == 0:
                # keys = tuple(text)
                properties = constructWordFieldHeader(properties_words, row)
                continue
            text = [unicodedata.normalize("NFKD", cell.text) for cell in row.cells]            
            type = getRowType(row, properties_words, nameField)
            # 最外边的那个节点
            if(type == "Order"):
                node = Node(properties, text, nameField)
                parentChain = [node.name]                
                yield node, parentChain
            elif(type == "Family"):
                node = Node(properties, text, nameField)
                if(len(parentChain) > 1):
                    parentChain = parentChain[:-1]
                parentChain.append(node.name)
                yield node, parentChain                                                  
            else:
                node = Node(properties, text, nameField)
                # currentFamily.addChild(node)
            


def scanWebPage(url):
    html = urllib.request.urlopen(url).read()
    soup = Soup(html, "html.parser")
    soup.select("a")
    return soup

def GetTag(element):
        return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))

def GetParagraphText(paragraph):    
    text = ''
    runCount = 0
    for child in paragraph._p:
        tag = GetTag(child)
        if tag == "w:r":
            text += paragraph.runs[runCount].text
            runCount += 1
        if tag == "w:hyperlink":
            for subChild in child:
                if GetTag(subChild) == "w:r":
                    text += subChild.text
    return text
# ocument.part.rels[rid]
# rid = child.attrib.items()[0][1]
def GetLinkUrl(paragraph):
    res = None
    for child in paragraph._p:
        if GetTag(child) == "w:hyperlink":
            rid = child.attrib.items()[0][1]
            if(rid not in paragraph.part.rels.keys()):
                break
            res = paragraph.part.rels[rid].target_ref        
    return res

# 判断word当前行是哪种类型
def getRowType(row, keys, nameField="中文名"):
    # check if row is merged
    _tcList = [cell._tc for cell in row.cells]
    # check _tcList item is all same
    # if true then row is merged, means it is a 苔藓植物 Bryophytes line
    if len(set(_tcList)) == 1:
        return "Order"
    
    # is not merged, check if it is "科" a 白发藓科 line
    text = (unicodedata.normalize("NFKD", cell.text) for cell in row.cells) 
    row_data = dict(zip(keys, text))
    if(str.strip(row_data[nameField])[-1] == "科"):
        return "Family"
    

    return "species"

def constructWordFieldHeader(properties_word, row):
    res = {}
    text = [unicodedata.normalize("NFKD", cell.text) for cell in row.cells]
    for i, word in enumerate(text):
        if word in properties_word:
            res[word] = i

    return res

if __name__ == '__main__':
    file_path = "./国家重点保护野生植物名录.docx"
    result_path = "./result_plant_word"
    document = Document(file_path)        

    properties_words = ["中文名", "学名", "保护级别", "备注"]
    nameField = "中文名"
    
    specTree = Node("植物", properties_words)
   

    utils.mkdir(result_path)
    for node in createGenerator(document, properties_words, nameField):
        print(node.name)

    # for tb in  document.tables:
    #     for i, row in enumerate(tb.rows):
                        
    #         if i == 0:
    #             # keys = tuple(text)
    #             properties = constructWordFieldHeader(properties_words, row)
    #             continue

    #         text = [unicodedata.normalize("NFKD", cell.text) for cell in row.cells]            
    #         # row_data = dict(zip(keys, text))
    #         type = getRowType(row, properties_words, nameField)
    #         # 最外边的那个节点
    #         if(type == "Order"):
    #             node = Node(properties, text, nameField)
    #             specTree.addChild(node)
    #             currentGroup = node
    #         elif(type == "Family"):
    #             node = Node(properties, text, nameField)
    #             currentGroup.addChild(node)
    #             currentFamily = node
    #         else:
    #             node = Node(properties, text, nameField)
    #             currentFamily.addChild(node)

        #     # data.append(row_data)
        #     for cell in row.cells:
        #         for p in cell.paragraphs:
        #             a = GetLinkUrl(p)
        #             print(a)
        #             print(p.text)
        # pass